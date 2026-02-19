#!/usr/bin/env python3
"""
Simple Resume Text Extractor
Extracts text from PDF and DOCX files using pdfplumber and python-docx
No AI - just raw text extraction
"""

import sys
import json
import io
import base64
import re

def extract_from_pdf(file_bytes):
    """Extract text from PDF using pdfplumber"""
    try:
        import pdfplumber
        pdf_file = io.BytesIO(file_bytes)
        text = ""
        
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"PDF extraction error: {str(e)}")

def extract_from_docx(file_bytes):
    """Extract text from DOCX using python-docx"""
    try:
        import docx
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        raise Exception(f"DOCX extraction error: {str(e)}")

def extract_text(file_base64, file_type):
    """
    Extract text from resume file
    
    Args:
        file_base64: Base64 encoded file content
        file_type: File extension (pdf, docx)
    
    Returns:
        Extracted text as string
    """
    try:
        # Remove data URL prefix if present
        if ',' in file_base64:
            file_base64 = file_base64.split(',')[1]
        
        # Decode base64
        file_bytes = base64.b64decode(file_base64)
        
        # Extract based on file type
        file_type = file_type.lower().replace('.', '')
        
        if file_type == 'pdf':
            text = extract_from_pdf(file_bytes)
        elif file_type == 'docx':
            text = extract_from_docx(file_bytes)
        else:
            raise Exception(f"Unsupported file type: {file_type}")
        
        if len(text.strip()) < 50:
            raise Exception("Extracted text is too short. File may be empty or corrupted.")
        
        return text
    
    except Exception as e:
        raise Exception(f"Text extraction failed: {str(e)}")

def parse_contact_info(text):
    """Extract contact information using regex"""
    # Email
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(email_pattern, text)
    email = emails[0] if emails else None
    
    # Phone
    phone_pattern = r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}'
    phones = re.findall(phone_pattern, text)
    phone = phones[0] if phones else None
    
    # LinkedIn
    linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
    linkedins = re.findall(linkedin_pattern, text, re.IGNORECASE)
    linkedin = linkedins[0] if linkedins else None
    
    # Name (first line, assuming it's the name)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    name = lines[0] if lines else None
    
    return {
        'name': name,
        'email': email,
        'phone': phone,
        'linkedin': linkedin
    }

def parse_skills(text):
    """Extract skills from text"""
    # Common technical skills list
    skill_keywords = [
        'JavaScript', 'TypeScript', 'Python', 'Java', 'C++', 'C#', 'Ruby', 'PHP', 'Go', 'Rust',
        'React', 'Angular', 'Vue', 'Node.js', 'Express', 'Django', 'Flask', 'Spring', 'Laravel',
        'MongoDB', 'MySQL', 'PostgreSQL', 'Redis', 'SQL', 'NoSQL', 'Elasticsearch',
        'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Jenkins', 'CI/CD',
        'Git', 'GitHub', 'GitLab', 'Jira', 'Agile', 'Scrum',
        'HTML', 'CSS', 'SASS', 'Bootstrap', 'Tailwind',
        'REST', 'API', 'GraphQL', 'Microservices',
        'Machine Learning', 'AI', 'Deep Learning', 'TensorFlow', 'PyTorch',
        'Linux', 'Unix', 'Bash', 'Shell',
        'Testing', 'Jest', 'Mocha', 'Selenium', 'Cypress'
    ]
    
    found_skills = []
    text_lower = text.lower()
    
    for skill in skill_keywords:
        if skill.lower() in text_lower:
            found_skills.append(skill)
    
    return found_skills

def extract_bio(text):
    """Extract summary/bio section from resume"""
    # Common bio section headers
    bio_headers = ['summary', 'profile', 'objective', 'about', 'bio']
    
    lines = text.split('\n')
    bio_text = ""
    capturing = False
    
    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        
        # Check if this line is a bio header
        if any(header in line_lower for header in bio_headers):
            capturing = True
            continue
        
        # Stop capturing at next section
        if capturing and line.isupper() and len(line) > 5:
            break
        
        # Capture bio lines
        if capturing and line.strip():
            bio_text += line.strip() + " "
            
            # Limit bio to 3-4 lines
            if len(bio_text) > 300:
                break
    
    # If no bio section found, use first few lines after name
    if not bio_text and len(lines) > 2:
        bio_text = " ".join(lines[2:5])
    
    return bio_text.strip()[:500]  # Max 500 chars

def main():
    """Main entry point"""
    try:
        # Read input from stdin or argument
        if len(sys.argv) > 1:
            input_json = sys.argv[1]
        else:
            input_json = sys.stdin.read()
        
        # Parse input
        input_data = json.loads(input_json)
        file_base64 = input_data.get('fileBase64')
        file_type = input_data.get('fileType')
        
        if not file_base64 or not file_type:
            raise Exception("Missing required fields: fileBase64 and fileType")
        
        # Extract text
        text = extract_text(file_base64, file_type)
        
        # Parse information
        contact = parse_contact_info(text)
        skills = parse_skills(text)
        bio = extract_bio(text)
        
        # Return result
        result = {
            "success": True,
            "text": text,
            "length": len(text),
            "contact": contact,
            "skills": skills,
            "bio": bio
        }
        
        print(json.dumps(result))
        sys.exit(0)
    
    except Exception as e:
        error_result = {
            "success": False,
            "error": str(e)
        }
        print(json.dumps(error_result))
        sys.exit(1)

if __name__ == "__main__":
    main()
